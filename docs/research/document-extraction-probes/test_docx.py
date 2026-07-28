"""Run every candidate DOCX extractor against resume_probe.docx and print which of the
seven planted markers each one recovers.

Apache Tika is optional: set TIKA_JAR to a tika-app JAR to include it, e.g.
  curl -sLo /tmp/tika-app.jar \\
    https://repo1.maven.org/maven2/org/apache/tika/tika-app/3.2.3/tika-app-3.2.3.jar
  TIKA_JAR=/tmp/tika-app.jar python3 test_docx.py

Run make_docx.py first.
"""

import os
import subprocess
from pathlib import Path

PATH = Path(__file__).parent / "out" / "resume_probe.docx"

MARKERS = [
    "BODY_PARAGRAPH_TEXT",
    "TABLE_LEFT_CELL",
    "HEADER_TEXT_PHONE_555",
    "FOOTER_TEXT_EMAIL",
    "LINK_ANCHOR_TEXT",
    "linkedin.com/in/realtarget",  # the URL itself, not merely the anchor text
    "TEXTBOX_SKILLS_PYTHON_SQL",
]


def python_docx_naive() -> str:
    """What a typical first implementation does: .paragraphs and nothing else."""
    import docx
    return "\n".join(p.text for p in docx.Document(str(PATH)).paragraphs)


def python_docx_thorough() -> str:
    """paragraphs + tables + headers/footers + hyperlink addresses."""
    import docx
    doc = docx.Document(str(PATH))
    out = [p.text for p in doc.paragraphs]
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                out.append(cell.text)
    for section in doc.sections:
        out += [p.text for p in section.header.paragraphs]
        out += [p.text for p in section.footer.paragraphs]
    for p in doc.paragraphs:
        for link in p.hyperlinks:
            out.append(link.address)
    return "\n".join(out)


def docx2python_text() -> str:
    from docx2python import docx2python
    with docx2python(str(PATH)) as d:
        return d.text


def docx2txt_process() -> str:
    import docx2txt
    return docx2txt.process(str(PATH))


def mammoth_raw() -> str:
    import mammoth
    with open(PATH, "rb") as f:
        return mammoth.extract_raw_text(f).value


def mammoth_html() -> str:
    import mammoth
    with open(PATH, "rb") as f:
        return mammoth.convert_to_html(f).value


def tika(mode: str):
    """Return a callable running `java -jar tika-app.jar --text|--html`."""
    def run() -> str:
        jar = os.environ.get("TIKA_JAR")
        if not jar or not Path(jar).exists():
            return "<<SKIPPED: set TIKA_JAR to a tika-app jar>>"
        proc = subprocess.run(
            ["java", "-jar", jar, mode, str(PATH)],
            capture_output=True, text=True, timeout=180,
        )
        return proc.stdout
    return run


EXTRACTORS = [
    ("python-docx .paragraphs only", python_docx_naive),
    ("python-docx thorough", python_docx_thorough),
    ("docx2python .text", docx2python_text),
    ("docx2txt.process", docx2txt_process),
    ("mammoth extract_raw_text", mammoth_raw),
    ("mammoth convert_to_html", mammoth_html),
    ("Apache Tika --text", tika("--text")),
    ("Apache Tika --html", tika("--html")),
]


def main() -> None:
    if not PATH.exists():
        raise SystemExit(f"missing {PATH} - run make_docx.py first")

    results: dict[str, str] = {}
    for name, fn in EXTRACTORS:
        try:
            results[name] = fn() or ""
        except Exception as e:
            results[name] = f"<<ERROR: {type(e).__name__}: {e}>>"

    width = max(len(k) for k in results) + 2
    short = ["body", "table", "header", "footer", "anchor", "link URL", "text box"]
    print("TOOL".ljust(width) + "".join(s.ljust(11) for s in short))
    print("-" * (width + 11 * len(short)))
    for name, text in results.items():
        row = name.ljust(width)
        for m in MARKERS:
            row += ("FOUND" if m in text else "MISS").ljust(11)
        print(row)

    print()
    for name, text in results.items():
        if text.startswith("<<"):
            print(f"{name}: {text.splitlines()[0]}")


if __name__ == "__main__":
    main()
