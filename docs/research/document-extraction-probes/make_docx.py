"""Build resume_probe.docx, exercising the DOCX features real resumes actually use.

Seven planted markers, each testing a distinct extraction path:
  BODY_PARAGRAPH_TEXT        plain body paragraph
  TABLE_LEFT/RIGHT_CELL      2-column table (designed resumes often use tables as layout)
  HEADER_TEXT_PHONE_555      section header (contact details frequently live here)
  FOOTER_TEXT_EMAIL          section footer
  LINK_ANCHOR_TEXT           w:hyperlink anchor text ...
  linkedin.com/in/realtarget  ... whose target URL differs from the anchor
  TEXTBOX_SKILLS_PYTHON_SQL  floating text box: mc:AlternateContent > wps:txbx > w:txbxContent
"""

from pathlib import Path

import docx
from docx.oxml import parse_xml

OUT = Path(__file__).parent / "out"
OUT.mkdir(exist_ok=True)
PATH = OUT / "resume_probe.docx"

HYPERLINK_REL = (
    "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink"
)

TEXTBOX_XML = """<w:p xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main"
 xmlns:mc="http://schemas.openxmlformats.org/markup-compatibility/2006"
 xmlns:wps="http://schemas.microsoft.com/office/word/2010/wordprocessingShape"
 xmlns:a="http://schemas.openxmlformats.org/drawingml/2006/main"
 xmlns:wp="http://schemas.openxmlformats.org/drawingml/2006/wordprocessingDrawing"
 xmlns:v="urn:schemas-microsoft-com:vml"
 xmlns:wp14="http://schemas.microsoft.com/office/word/2010/wordprocessingDrawing">
 <w:r><mc:AlternateContent>
  <mc:Choice Requires="wps"><w:drawing>
   <wp:anchor distT="0" distB="0" distL="0" distR="0" simplePos="0" relativeHeight="251"
              behindDoc="0" locked="0" layoutInCell="1" allowOverlap="1">
    <wp:simplePos x="0" y="0"/>
    <wp:positionH relativeFrom="column"><wp:posOffset>0</wp:posOffset></wp:positionH>
    <wp:positionV relativeFrom="paragraph"><wp:posOffset>0</wp:posOffset></wp:positionV>
    <wp:extent cx="2000000" cy="1000000"/>
    <wp:effectExtent l="0" t="0" r="0" b="0"/>
    <wp:wrapNone/>
    <wp:docPr id="1" name="TextBox 1"/>
    <a:graphic><a:graphicData uri="http://schemas.microsoft.com/office/word/2010/wordprocessingShape">
     <wps:wsp><wps:cNvSpPr txBox="1"/>
      <wps:spPr><a:xfrm><a:off x="0" y="0"/><a:ext cx="2000000" cy="1000000"/></a:xfrm>
       <a:prstGeom prst="rect"><a:avLst/></a:prstGeom></wps:spPr>
      <wps:txbx><w:txbxContent>
        <w:p><w:r><w:t>TEXTBOX_SKILLS_PYTHON_SQL</w:t></w:r></w:p>
      </w:txbxContent></wps:txbx>
      <wps:bodyPr rot="0" vert="horz" wrap="square"/>
     </wps:wsp></a:graphicData></a:graphic>
   </wp:anchor></w:drawing></mc:Choice>
  <mc:Fallback><w:pict><v:rect style="width:150pt;height:75pt"/></w:pict></mc:Fallback>
 </mc:AlternateContent></w:r>
</w:p>"""


def main() -> None:
    d = docx.Document()
    d.add_paragraph("BODY_PARAGRAPH_TEXT")

    table = d.add_table(rows=1, cols=2)
    table.cell(0, 0).text = "TABLE_LEFT_CELL"
    table.cell(0, 1).text = "TABLE_RIGHT_CELL"

    section = d.sections[0]
    section.header.paragraphs[0].text = "HEADER_TEXT_PHONE_555"
    section.footer.paragraphs[0].text = "FOOTER_TEXT_EMAIL"

    para = d.add_paragraph()
    r_id = para.part.relate_to(
        "https://linkedin.com/in/realtarget", HYPERLINK_REL, is_external=True
    )
    para._p.append(parse_xml(
        '<w:hyperlink '
        'xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main" '
        'xmlns:r="http://schemas.openxmlformats.org/officeDocument/2006/relationships" '
        f'r:id="{r_id}"><w:r><w:t>LINK_ANCHOR_TEXT</w:t></w:r></w:hyperlink>'
    ))

    d.element.body.insert(0, parse_xml(TEXTBOX_XML))
    d.save(str(PATH))
    print(f"wrote {PATH}")


if __name__ == "__main__":
    main()
